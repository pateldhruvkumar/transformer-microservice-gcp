"""
generate_guide_pdf.py — Generates the professor user guide as a Word document.
Run from the repo root:
    python3 scripts/generate_guide_pdf.py
Output: Shakespeare_Text_Generator_Guide.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SERVICE_URL = "https://shakespeare-api-925816807412.us-central1.run.app/ui"

ORANGE = RGBColor(0xFF, 0x89, 0x06)
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
GRAY   = RGBColor(0x78, 0x78, 0x96)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = ORANGE
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.runs[0].font.size = Pt(10.5)
    return p


def add_table(doc, headers, rows, col_widths_in):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Inches(col_widths_in[i])
        set_cell_bg(cell, "1A1A2E")
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = ORANGE
        run.font.size = Pt(9)

    # data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = "F0F0F5" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.width = Inches(col_widths_in[ci])
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)

    doc.add_paragraph()


def build():
    doc = Document()

    # page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── title block ──────────────────────────────────────────────────────────
    title = doc.add_heading("Shakespeare Text Generator - User Guide", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = ORANGE

    sub = doc.add_paragraph("Transformer Language Model  |  FastAPI Microservice  |  Google Cloud Run")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = GRAY
    sub.runs[0].font.size = Pt(10)

    course = doc.add_paragraph("EAI6010 - Assignment 5")
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    course.runs[0].font.color.rgb = GRAY
    course.runs[0].font.size = Pt(9)
    course.runs[0].italic = True

    doc.add_paragraph()

    # ── live url ─────────────────────────────────────────────────────────────
    add_heading(doc, "Live Service URL", level=1)
    add_body(doc, "Open this link in any browser to access the UI:")

    p = doc.add_paragraph()
    run = p.add_run(SERVICE_URL)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

    doc.add_paragraph()

    # ── what it does ─────────────────────────────────────────────────────────
    add_heading(doc, "What It Does", level=1)
    add_body(doc,
        "This service runs a PyTorch Transformer language model trained on the Tiny Shakespeare "
        "dataset (~1.1 MB of Shakespeare's plays). Given a short seed prompt, the model generates "
        "a continuation in Shakespearean style using autoregressive token prediction."
    )
    doc.add_paragraph()

    # ── step 1 ───────────────────────────────────────────────────────────────
    add_heading(doc, "Step 1 - Enter a Prompt", level=1)
    add_body(doc, "Type any seed text in the Prompt box. The model will continue from where you leave off.")
    add_body(doc, "Suggested prompts to try:")
    for prompt in [
        "to be or not to be",
        "the king shall rise",
        "romeo , romeo , wherefore art thou",
        "my lord , i have",
        "thou art a fool",
    ]:
        add_bullet(doc, prompt)
    doc.add_paragraph()

    # ── step 2 ───────────────────────────────────────────────────────────────
    add_heading(doc, "Step 2 - Pick a Generation Mode", level=1)
    add_body(doc, "Click one of the three preset buttons - they auto-fill Temperature and Top-K:")
    add_table(doc,
        headers=["Mode", "Temperature", "Top-K", "Behaviour", "Best For"],
        rows=[
            ["Greedy",   "1.0", "0",  "Always picks most likely word. Deterministic.", "Safest prediction"],
            ["Balanced", "0.8", "10", "Samples from top 10 candidates. Coherent & varied.", "General use"],
            ["Creative", "1.3", "20", "Wider sampling. More surprising output.", "Exploring range"],
        ],
        col_widths_in=[0.9, 0.9, 0.6, 2.8, 1.2],
    )

    # ── step 3 ───────────────────────────────────────────────────────────────
    add_heading(doc, "Step 3 - Adjust Parameters (Optional)", level=1)
    add_body(doc, "Fine-tune the generation by editing the numeric fields directly:")
    add_table(doc,
        headers=["Parameter", "Default", "Range", "What It Controls"],
        rows=[
            ["Max Tokens",  "50",  "1 - 200",   "How many words to generate after the prompt"],
            ["Temperature", "0.8", "0.1 - 2.0", "Lower = conservative, Higher = creative/random"],
            ["Top-K",       "10",  "0 - 50",    "Candidate pool size. 0 means greedy (no sampling)"],
        ],
        col_widths_in=[1.2, 0.8, 0.9, 3.5],
    )

    # ── step 4 ───────────────────────────────────────────────────────────────
    add_heading(doc, "Step 4 - Click Generate", level=1)
    add_bullet(doc, "The orange text is your original prompt.")
    add_bullet(doc, "The white text is what the model generated.")
    add_bullet(doc, "Below the output: tokens generated + parameters used.")
    doc.add_paragraph()

    # ── example response ─────────────────────────────────────────────────────
    add_heading(doc, "Example API Response", level=1)
    example = (
        '{\n'
        '  "prompt": "to be or not to be",\n'
        '  "generated_text": "to be or not to be his son , sir , but i will not ...",\n'
        '  "parameters": { "max_tokens": 50, "temperature": 0.8, "top_k": 10 },\n'
        '  "tokens_generated": 50\n'
        '}'
    )
    p = doc.add_paragraph()
    run = p.add_run(example)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    doc.add_paragraph()

    # ── important notes ───────────────────────────────────────────────────────
    add_heading(doc, "Important Notes", level=1)
    add_bullet(doc,
        "Cold Start: The first request after inactivity may take 5-10 seconds. "
        "Cloud Run scales to zero when idle and reloads the model (~22 MB) on the first request."
    )
    add_bullet(doc, "Greedy mode tends to repeat phrases - this is a known limitation of greedy decoding.")
    add_bullet(doc, "Balanced or Creative mode produces more natural-looking Shakespearean output.")
    add_bullet(doc, "Increase Max Tokens (e.g. 100-150) for longer passages.")
    add_bullet(doc, "Full REST API docs available at: /docs endpoint.")
    doc.add_paragraph()

    # ── api endpoints ─────────────────────────────────────────────────────────
    add_heading(doc, "API Endpoints (for reference)", level=1)
    add_table(doc,
        headers=["Method", "Endpoint", "Description"],
        rows=[
            ["GET",  "/",         "Service metadata and status"],
            ["GET",  "/health",   "Health check - returns {\"status\": \"healthy\"}"],
            ["GET",  "/ui",       "Web user interface"],
            ["POST", "/generate", "Generate text from a prompt"],
            ["GET",  "/docs",     "Interactive Swagger API documentation"],
        ],
        col_widths_in=[0.7, 1.1, 4.6],
    )

    out = "Shakespeare_Text_Generator_Guide.docx"
    doc.save(out)
    print(f"Document saved -> {out}")


if __name__ == "__main__":
    build()
